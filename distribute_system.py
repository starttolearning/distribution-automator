# Automatic Distribute Stuff Workout System
import prettytable as pt
import random
from xlwt import Workbook, easyxf
from xlrd import open_workbook
import xlrd
import xlwt

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
import datetime


# Define status code for job control
# distribute in average: -1
# distribute in specific times when drop to 0
# distribute maxiumn: -2
EVNELY_DISTRIBUTE = -1
OUT_OF_DISTRIBUTE = 0
MAXIUMN_DISTRIBUTE = -2

# Register job table
DATA_FROM = 'job_dist.xlsx'
WEEK_DISTRIBUTION_DOCX = 'meizhouxx.docx'
CLASS_NAMES = ['07', '08', '09']
START_DATE = '2020-09-01'
TOTAL_WEEKS = 16


"""
This function return a tupe which contians all the
rigister teacher distribution and every teacher's
job control settings
"""


def get_teacher_expect_distribute(data_from, sheet_name):
    sh = open_workbook(data_from).sheet_by_name(sheet_name)
    nrows, ncols = sh.nrows, sh.ncols
    datas = {}
    job_controls = {}
    teacher_id_name = {}

    for x in range(1, nrows):
        data = sh.row_values(x)
        if data[2] not in datas.keys():
            datas[data[2]] = []
            job_controls[data[2]] = {}

        datas[data[2]].append(data)
        job_controls[data[2]][data[0]] = int(data[3])
        teacher_id_name[data[0]] = str(data[1])

    return (datas, job_controls, teacher_id_name)

# Hold all the teacher_datas_and_job and job_controls information from the spreadsheet


def get_three_class_distribution_information():
    three_class_datas = {}
    three_class_job_controls = {}
    three_class_teacher_id_names = {}
    three_class_count = {}

    for sheet_name in CLASS_NAMES:
        datas, job_controls, teacher_id_name = get_teacher_expect_distribute(
            DATA_FROM, sheet_name)
        three_class_datas[sheet_name] = datas
        three_class_job_controls[sheet_name] = job_controls
        three_class_teacher_id_names[sheet_name] = teacher_id_name
        three_class_count[sheet_name] = len(datas)
    return (three_class_datas,
            three_class_job_controls,
            three_class_teacher_id_names,
            three_class_count)


"""
Using job controls dictionary to calculate next piority class

:Parameters
job_control     dict
"""


def get_class_piority(job_controls):
    job_piority = {}
    job_piority_temps = {}
    for key, jobs in job_controls.items():
        total = 0
        for job in jobs.values():
            total += job
        job_piority_temps[key] = total / len(jobs)

    job_piority_temps = sorted(
        job_piority_temps.items(), key=lambda d: d[1], reverse=True)

    for job in job_piority_temps:
        job_piority[job[0]] = job[1]
    return job_piority


# This function get a week of job distributions


def get_a_week_of_distribute(datas, job_controls, class_count):
    week_of_distribute = {}  # This hold a week of  distribute
    expect_list_in_key = []  # Using this to hold the expect keys
    job_controls_bak = job_controls.copy()

    # Loop through all the item in the data set to find a
    # valide list
    class_piority = get_class_piority(job_controls_bak)

    for key in class_piority.keys():
        # print(key, value)
        job_control = job_controls_bak[key]
        times = 100
        while(True):
            times = times-1
            if times < 0:
                break

            teacher_id, name, class_no, control = random.choice(datas[key])
            # print(teacher_id, name, class_no, control)
            if teacher_id not in expect_list_in_key:
                status = job_control.get(teacher_id, -3)
                if status > 0:
                    job_control.update(
                        {teacher_id: job_control.get(teacher_id) - 1})
                    expect_list_in_key.append(teacher_id)
                    week_of_distribute[key] = str(name)
                    break
                elif status == OUT_OF_DISTRIBUTE:
                    continue
                elif status == EVNELY_DISTRIBUTE:
                    expect_list_in_key.append(teacher_id)
                    week_of_distribute[key] = str(name)
                    break
    # print(len(week_of_distribute))
    if len(week_of_distribute) != class_count:
        return []

    week_of_distribute = sorted(
        week_of_distribute.items(), key=lambda d: d[0])

    week_of_distribute_temp = []
    for item in week_of_distribute:
        week_of_distribute_temp.append(item[1])

    job_controls.update(job_controls_bak)
    return week_of_distribute_temp


def prettytable_output(distribution_sets):
    tb = pt.PrettyTable()
    tb.field_names = distribution_sets[0]
    for distribution_set in distribution_sets[1:]:
        tb.add_row(distribution_set)
    print(tb)

# print(datas)


"""
Get distribution sets

"""


def get_distribution_sets(datas, job_controls, class_count):
    distribution_sets = []
    distribution_sets_temps = []
    field_names = ["周 次"] + [str(x+1)+" 班" for x in range(class_count)]
    for i in range(TOTAL_WEEKS):
        ready_to_add = get_a_week_of_distribute(
            datas, job_controls, class_count)
        if ready_to_add:
            distribution_sets_temps.append(ready_to_add)

    distribution_sets_temps_length = len(distribution_sets_temps)
    for x in range(distribution_sets_temps_length):
        level_one = 10
        while(len(distribution_sets_temps)):
            level_one -= 1
            distribution_set_temp = random.choice(distribution_sets_temps)
            index_of_distribution_set = distribution_sets_temps.index(
                distribution_set_temp)
            if distribution_sets and distribution_sets[len(distribution_sets)-1][1] != distribution_set_temp[1]:
                distribution_sets.append(
                    ["第"+str(x+1)+"周"] + distribution_set_temp)
                distribution_sets_temps.pop(index_of_distribution_set)
                break
            elif len(distribution_sets) == 0:
                distribution_sets.append(
                    ["第"+str(x+1)+"周"] + distribution_set_temp)
                distribution_sets_temps.pop(index_of_distribution_set)
                break
            if level_one < 0:
                distribution_sets.append(
                    ["第"+str(x+1)+"周"] + distribution_set_temp)
                distribution_sets_temps.pop(index_of_distribution_set)
                break
    distribution_sets.insert(0, field_names)
    return distribution_sets


"""
Write the distribution sets to a work book
"""


def write_distribution_to_workbook(workbook, distribution_sets, class_name, class_count):
    # Setting up for the workbook and worksheet
    distribution_worksheet = workbook.add_sheet(
        'distribution_sets_'+class_name)
    # Setting up for the title
    style = xlwt.XFStyle()
    font = xlwt.Font()
    font.bold = True
    alignment = xlwt.Alignment()
    alignment.horz = xlwt.Alignment.HORZ_CENTER
    alignment.vert = xlwt.Alignment.VERT_CENTER
    style.font = font
    style.alignment = alignment
    distribution_worksheet.write_merge(
        0, 0, 0, class_count, "Distribution Sets", style)

    # Setting up for the fields row

    for distribution_set in distribution_sets:
        row_of_distribution = distribution_sets.index(distribution_set) + 1
        is_field_row = True if row_of_distribution == 1 else False

        for distribution_item in distribution_set:
            col_of_distribution = distribution_set.index(distribution_item)
            if is_field_row:
                distribution_worksheet.write(
                    row_of_distribution, col_of_distribution, distribution_item, style)
            else:
                distribution_worksheet.write(
                    row_of_distribution, col_of_distribution, distribution_item)
    workbook.save('distributions.xls')


"""
Write the distribution sets to word, make each set a page
"""


def build_distribution_sets_for_word_document(distribution_sets):
    distribution_sets_by_class = []
    for distribution_set in distribution_sets[1:]:
        distribution_sets_by_class.append(distribution_set[1:])
    return distribution_sets_by_class


def write_distribution_to_word_document(three_class_distribution_sets):
    document = Document()
    max_cols = 9
    max_rows = 18
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles['Title'].font.name = 'MFShangYa_Noncommercial-Regular'
    document.styles['Normal']._element.rPr.rFonts.set(
        qn('w:eastAsia'), u'MFShangYa_Noncommercial-Regular')
    document.styles['Subtitle'].font.name = u'微软雅黑'
    document.styles['Subtitle']._element.rPr.rFonts.set(
        qn('w:eastAsia'), u'微软雅黑')
    paragraph = document.add_paragraph(
        '2020-2021秋季学期盐边县红旭青少年俱乐部（渔门中学）', style='Title')
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = document.add_paragraph('活动项目指导人员考情表', style='Subtitle')
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()
    """
    Table header: 第X周（XX月XX日-XX月XX日）
    Table fields: 序号、星期二、考情情况、星期三、考情情况、星期四、考情情况
    序号从1开始，代表班级
    最后一行是考情
    """
    # Table header:第X周（XX月XX日-XX月XX日）
    # heading_cells = table.rows[0].cells
    table_fields = ['序号', '星期二', '考情情况', '星期三',
                    '考情情况', '星期四', '考情情况', '星期五', '考情情况']
    x1, x2, x3 = three_class_distribution_sets.values()
    total_week = min(len(x1), len(x2), len(x3))

    start_date = datetime.datetime.strptime(START_DATE, '%Y-%m-%d')

    for week_identifer in range(total_week):
        print("正在写入：第"+str(week_identifer+1)+"周的排列")
        table = document.add_table(rows=max_rows, cols=max_cols)
        table.style = 'Table Grid'
        table.autofit = True
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.rows[0].cells[0].merge(table.rows[0].cells[8])

        # date compute start with Tuesday, end with Friday
        for n_days in range(7):
            start_date += datetime.timedelta(days=1)
            if datetime.datetime.weekday(start_date) == 1:
                break

        # Write the first row
        title_row_output = "第 {} 周 ({}-{})".format(week_identifer+1, start_date.strftime(
            '%Y-%m-%d'), (start_date+datetime.timedelta(days=3)).strftime('%Y-%m-%d'))
        table.rows[0].cells[0].text = title_row_output
        table.rows[0].cells[0].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        # write the seconde row
        for col_index, item in enumerate(table_fields):
            table.rows[1].cells[col_index].text = item

        index = 0
        x1_length, x2_length, x3_length = len(x1[week_identifer]), len(
            x2[week_identifer]), len(x3[week_identifer])
        x1_x2_x3_max_length = max(x1_length, x2_length, x3_length)

        if x1_x2_x3_max_length-x1_length > 0:
            for x in range(x1_x2_x3_max_length-x1_length):
                x1[week_identifer].append([])
        if x1_x2_x3_max_length-x2_length > 0:
            for x in range(x1_x2_x3_max_length-x2_length):
                x2[week_identifer].append([])
        if x1_x2_x3_max_length-x3_length > 0:
            for x in range(x1_x2_x3_max_length-x3_length):
                x3[week_identifer].append([])

        for class_one, class_two, class_three in zip(x1[week_identifer], x2[week_identifer], x3[week_identifer]):
            table.cell(index+2, 0).text = str(index+1)
            table.cell(index+2, 0).width = Cm(1)

            table.cell(index+2, 1).text = class_one
            table.cell(index+2, 1).width = Cm(2)
            table.cell(index+2, 2).width = Cm(1.6)

            table.cell(index+2, 3).text = class_two
            table.cell(index+2, 3).width = Cm(2)
            table.cell(index+2, 4).width = Cm(1.6)

            table.cell(index+2, 5).text = class_three
            table.cell(index+2, 5).width = Cm(2)
            table.cell(index+2, 6).width = Cm(1.6)

            table.cell(index+2, 7).width = Cm(2)

            index += 1

        for row in table.rows:
            row.height_rule = WD_ROW_HEIGHT.EXACTLY
            row.height = Cm(0.94)
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(14)

        table.rows[max_rows-1].cells[0].text = '考情情况'
        table.rows[max_rows-1].height = Cm(1.5)
        table.rows[1].height = Cm(1.6)
        if week_identifer != total_week-1:
            document.add_page_break()
        print("写入成功：第"+str(week_identifer+1)+"周的排列！！")
    document.save(WEEK_DISTRIBUTION_DOCX)


three_class_datas, three_class_job_controls, three_class_teacher_id_names, three_class_count = get_three_class_distribution_information()
three_class_distribution_sets = {}
workbook = Workbook(encoding='utf-8')
print("正在写入，请稍后……")
for class_name in CLASS_NAMES:

    distribution_sets = get_distribution_sets(
        three_class_datas[class_name], three_class_job_controls[class_name], three_class_count[class_name])
    three_class_distribution_sets[class_name] = build_distribution_sets_for_word_document(
        distribution_sets)
    write_distribution_to_workbook(workbook,
                                   distribution_sets, class_name, three_class_count[class_name])
# print(three_class_distribution_sets)
# prettytable_output(distribution_sets)
# write_distribution_to_workbook(distribution_sets)
write_distribution_to_word_document(three_class_distribution_sets)
print("成功！！请查看排列总表："+DATA_FROM+", 每周情况表："+WEEK_DISTRIBUTION_DOCX)
